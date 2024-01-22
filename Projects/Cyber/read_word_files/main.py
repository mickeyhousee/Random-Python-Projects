from docx import Document
import os

def search_word(doc_path, word):
    try:
        # Open the Word document using python-docx
        doc = Document(doc_path)

        # Search for the word in paragraphs
        for paragraph in doc.paragraphs:
            if word.lower() in paragraph.text.lower():
                return True

        # Search for the word in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if word.lower() in cell.text.lower():
                        return True

    except Exception as e:
        # Handle the exception (e.g., print an error message)
        print(f"Error processing file {doc_path}: {e}")

    return False

def find_and_show_file_paths(word):
    found_files = []  # List to store paths of all files containing the word
    
    # Walk through the file system starting from the root of C:
    for root, dirs, files in os.walk('C:\\'):
        for file in files:
            # Check if the file is a Word document
            if file.endswith(".docx"):
                # Construct the full path of the file
                file_path = os.path.join(root, file)

                # Search for the word in the document
                if search_word(file_path, word):
                    found_files.append(file_path)

    return found_files

word_to_search = 'banana'

found_files = find_and_show_file_paths(word_to_search)

if found_files:
    print(f'The word "{word_to_search}" was found in the following files:')
    for file_path in found_files:
        print(file_path)
else:
    print(f'The word "{word_to_search}" was not found in any files in the C: drive.')
